import docx

import changeDirectory
# import csv
# import pathlib

def WordAddTextAfterListMarking(documentName,WordToReplaceList, NewWordToPutList):
    
    if len(WordToReplaceList)==len(NewWordToPutList):
        doc = docx.Document(documentName)                     
        for paragraph in doc.paragraphs:
            for i in range(len(WordToReplaceList)): 
                if WordToReplaceList[i] in paragraph.text:
                    orig_text = paragraph.text
                    new_text = orig_text.replace(WordToReplaceList[i], NewWordToPutList[i])
                    paragraph.text = new_text

                    print("Document Changed")
                    
            
        doc.save(documentName)

        print("Done")
    else:
        print("the two lists are not the same size")



existingDocumentFile='TextWordReplaceFuntion.docx'

paragraph="this worked"

bookmark="BLABLA"


Folder="AutoFormFillerOutputFiles"

changeDirectory.ChangeStartupDirectory(Folder=Folder)


documentName=existingDocumentFile


WordToReplace=bookmark


NewWordToPut=paragraph


# document = Document(existingDocumentFile)

# WordAddTextAfterMarking(document, paragraph, bookmark)

WordToReplaceList=["BLABLA1","BLABLA2","BLABLA3"]

NewWordToPutList=["this worked1","this worked2","this worked3"]


WordAddTextAfterListMarking(documentName,WordToReplaceList, NewWordToPutList)