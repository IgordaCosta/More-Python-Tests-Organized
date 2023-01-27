from docx import Document

import changeDirectory



# def WordAddTextAfterMarking(document, paragraph, bookmark = '[BOOKMARK HERE]'):
    
#   for para in document.paragraphs:
#     if para.text == bookmark:
#       p = para._p
#       p.addnext(paragraph._p)


import docx
# import csv
# import pathlib

def WordAddTextAfterMarking2(documentName,WordToReplace, NewWordToPut):
    
    doc = docx.Document(documentName)               
    for paragraph in doc.paragraphs:
        if WordToReplace in paragraph.text:
            orig_text = paragraph.text
            new_text = orig_text.replace(WordToReplace, NewWordToPut)
            paragraph.text = new_text

            print("Document Changed")
            
    #     if "AMNT" in paragraph.text:
    #         orig_text = paragraph.text
    #         new_text = str.replace(orig_text, "AMNT", line["Amount"])
    #         paragraph.text = new_text
    #     if "ADDR" in paragraph.text:
    #         orig_text = paragraph.text
    #         new_text = str.replace(orig_text, "ADDR", line["Address"])
    #         paragraph.text = new_text
    # print("Saving: " + line["Date"] + ".docx")
    doc.save(documentName)

    print("Done")



existingDocumentFile='TextWordReplaceFuntion.docx'

paragraph="this worked"

bookmark="BLABLA"


Folder="AutoFormFillerOutputFiles"

changeDirectory.ChangeStartupDirectory(Folder=Folder)


documentName=existingDocumentFile


WordToReplace=bookmark


NewWordToPut=paragraph


# document = Document(existingDocumentFile)

# WordAddTextAfterMarking(document, paragraph, bookmark)


WordAddTextAfterMarking2(documentName,WordToReplace, NewWordToPut)


# document.save(existingDocumentFile)