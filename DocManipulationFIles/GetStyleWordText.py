from docx import Document, text


def GetStyleWordText(DocName):

    # oxml.text.font
    document = Document(docx = DocName)

    print(document.paragraphs[2].text)
    print(document.paragraphs[2].style.font)



    # print(document.paragraphs[0].text)
    # print(document.paragraphs[0].style)

    # for i in document.paragraphs[0]:
    #     print(i)

    
    
    # print(document.paragraphs[0].oxml.sz_val())
    # font = document.styles.font

    # print(font)




    # for y in range(len(document.paragraphs)):
    #     paragraph = document.paragraphs[y]

    #     print(paragraph)

    #     paragraph.
        

        # print(paragraph.style)

        # for x in range(len(paragraph.style)):
        #     print(x)

        # for x in range(len(paragraph.text)):
        #     letter = paragraph.text[x]

        #     print(letter)

           






Location = r'C:\Users\IgorDC\Desktop\PydocTest' +'\\'


DocName = Location + 'DifferentStylesTest.docx'



GetStyleWordText(DocName)
