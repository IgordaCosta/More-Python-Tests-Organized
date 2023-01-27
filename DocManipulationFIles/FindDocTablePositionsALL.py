import docx
from docx import Document
import pprint

import NoDuplicateList


filename='C:\\Users\\IgorDC\\Desktop\\PydocTest\\TestDocNoChange.docx'


splitter='XXXXXXXXXXX'

filename2='C:\\Users\\IgorDC\\Desktop\\PydocTest\\TestChanged.docx'


filename4='C:\\Users\\IgorDC\\Desktop\\PydocTest\\TestChanged4.docx'




def FindDocTablePositionsALL(filename):
    wordDoc = Document(filename)

    # add digit values from 1 to amount to end into all table variables
    # then get those values back, the result will be be variables with empty spots in the middle
    #  ex. 1,4,8,12,20,25 ...

    tableNumber=0
    for table in wordDoc.tables:

        rowNumber=0
        for row in table.rows:

            cellNumber=0
            id=0
            for cell in row.cells:

                wordDoc.tables[tableNumber].cell(rowNumber,cellNumber).text = str(id)

                id = id+1

                cellNumber=cellNumber+1

            rowNumber=rowNumber+1

        tableNumber=tableNumber+1

    PositionList=[]
    
    # gets the table row and cell position
    tableNumber=0
    for table in wordDoc.tables:

        rowNumber=0
       
        for row in table.rows:

            cellNumber=0
            
            tempList=[]
            for cell in row.cells:

                ValueText =wordDoc.tables[tableNumber].cell(rowNumber,cellNumber).text

                tempList.append(ValueText)

                cellNumber=cellNumber+1


            ValueTextListND = NoDuplicateList.NoDuplicateList(List=tempList)

            for i in range(len(ValueTextListND)):
                PositionList.append([tableNumber,rowNumber,int(ValueTextListND[i])])

            rowNumber=rowNumber+1

        tableNumber=tableNumber+1

    return PositionList





# PositionList = FindDocTablePositionsALL(filename)


# pprint.pprint(PositionList)


# print("Done")