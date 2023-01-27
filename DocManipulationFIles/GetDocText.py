import docx
from docx import Document

import IfWordInTextRemoveAll

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)



filename='C:\\Users\\IgorDC\\Desktop\\PydocTest\\TestDocNoChange.docx'

text=getText(filename=filename)

# print(text)

# print("text 1 above")

# print(len(text))

# print('initial size above')


splitterTable='XXXXXXXXXXX'

InputSplitter=' ' + splitterTable


filename2='C:\\Users\\IgorDC\\Desktop\\PydocTest\\TestChanged.docx'

text2=getText(filename=filename2)

# print(text2)

# print("text 2 above")

# print("initially text 1 == test 2 ?", text==text2 )

text3 = text2

text3 = text3.replace(InputSplitter, "")

# print("after text 1 == test 3 ?", text==text3 )

# print('after text2==text3 ?' , text2==text3)

OnlyLocationMaquerAddedCheck = text==text3

print('OnlyLocationMaquerAddedCheck :',OnlyLocationMaquerAddedCheck)

if OnlyLocationMaquerAddedCheck:

    word = InputSplitter

    stringUsed = text2

    ListOfNumLocations = IfWordInTextRemoveAll.IfWordInTextRemoveAll(word, stringUsed , ListOfWordPositions = [])

    print(ListOfNumLocations)








    # NumWord = FindNumberWordInString.FindNumberWordInString(word=word , stringUsed=stringUsed)

    # print('NumWord: ',NumWord)

    # WordLocationList = []
    # # for i in range(NumWord):
    # result = text2.find(word) 

    # print(result)

    # text3=text2.replace(word,"",1)

    # NumWord = FindNumberWordInString.FindNumberWordInString(word=word , stringUsed=text3)

    # print('NumWord: ',NumWord)



    




    

# Text2Lists=text2.split(splitterTable)

# sizes=[]
# AllNewLength=0
# for splits in Text2Lists:
#     size=len(splits)
#     sizes.append(size)
#     AllNewLength=AllNewLength+size

# print(sizes)

# print('sizes above')

# print(AllNewLength)

# print('AllNewLength above')



wordDoc = Document(filename2)


indexNumber=0
TableRowCell=[]
indexList=[]


# rowNumber=0
# cellNumber=0


# print(len(wordDoc.tables))

# print('len(wordDoc.tables)')



# add digit values from 1 to amount to end into all table variables
# then get those values back, the result will be be variables with empty spots in the middle
#  ex. 1,4,8,12,20,25 ...

tableNumber=0



for table in wordDoc.tables:

    # print('len(wordDoc.tables): '+ str(len(wordDoc.tables)))
    

    # print('tableNumber: ' + str(tableNumber))

    
    rowNumber=0
    for row in table.rows:

        # print("len(table.rows): " + str(len(table.rows)))
        
        # print('rowNumber: ' + str(rowNumber))

        
        

        cellNumber=0
        for cell in row.cells:

            # print('len(row.cells) :' + str(len(row.cells)))
            
            # print('cellNumber: ' + str(cellNumber))

            
            

            # print(cell.text)
            if cell.text == splitterTable:
                indexNumber=indexNumber+1
                TableRowCell.append([tableNumber,rowNumber,cellNumber])
                indexList.append(indexNumber)

            cellNumber=cellNumber+1

        rowNumber=rowNumber+1

    tableNumber=tableNumber+1


print(TableRowCell)

print('TableRowCell above')

# print(indexList)

# print('indexList above')



# filename3='C:\\Users\\IgorDC\\Desktop\\PydocTest\\TestChanged2.docx'



# wordDoc = docx.Document(filename)


# variable=['YYYYYYYY1','YYYYYYYY2','YYYYYYYY3','YYYYYYYY4','YYYYYYYY5','YYYYYYYY6','YYYYYYYY7','YYYYYYYY8', 'YYYYYYYY9']
 
# num = 0
# for num in range(len(TableRowCell)):
#     [t, r, c]= TableRowCell[num]
#     wordDoc.tables[t].cell(r,c).text = variable[num]
    

 
# wordDoc.save(filename3)
                






