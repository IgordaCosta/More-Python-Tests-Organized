import docx
from docx import Document

import pprint

import NoDuplicateList


filename='C:\\Users\\IgorDC\\Desktop\\PydocTest\\TestDocNoChange.docx'


splitter='XXXXXXXXXXX'

filename2='C:\\Users\\IgorDC\\Desktop\\PydocTest\\TestChanged.docx'


filename4='C:\\Users\\IgorDC\\Desktop\\PydocTest\\TestChanged4.docx'





# wordDoc = Document(filename4)


# indexNumber=0
# TableRowCell=[]
# indexList=[]


# print(len(wordDoc.tables))

# print('len(wordDoc.tables)')



# add digit values from 1 to amount to end into all table variables
# then get those values back, the result will be be variables with empty spots in the middle
#  ex. 1,4,8,12,20,25 ...
def GetTableDocPositions2(filename):
    wordDoc = Document(filename)
    # indexNumber=0
    # TableRowCell=[]
    # indexList=[]
    # ValueTextList=[]
    OneList=[]
    # NewValueList=[]
    tableNumber=0
    for table in wordDoc.tables:

        # print('len(wordDoc.tables): '+ str(len(wordDoc.tables)))
        

        # print('tableNumber: ' + str(tableNumber))

        # NewValueList.append([])
        rowNumber=0
        # ValueTextList.append([])
        for row in table.rows:

            # print("len(table.rows): " + str(len(table.rows)))
            
            # print('rowNumber: ' + str(rowNumber))

            
            

            cellNumber=0
            # ValueTextList[tableNumber].append([])
            tempList=[]
            for cell in row.cells:

                # print('len(row.cells) :' + str(len(row.cells)))
                
                # print('cellNumber: ' + str(cellNumber))


                ValueText =wordDoc.tables[tableNumber].cell(rowNumber,cellNumber).text

                

                

                # ValueTextList[tableNumber][rowNumber].append(ValueText)


                tempList.append(ValueText)

        

                
                
                cellNumber=cellNumber+1


            ValueTextListND = NoDuplicateList.NoDuplicateList(List=tempList)

            # NewValueList[tableNumber].append(ValueTextListND)

            for i in range(len(ValueTextListND)):
                OneList.append([tableNumber,rowNumber,ValueTextListND[i]])

            rowNumber=rowNumber+1


        tableNumber=tableNumber+1

    return OneList
    # pprint.pprint(NewValueList)

    # pprint.pprint(OneList)

# print(len(OneList))

# print('len(OneList) above')

# ni=0
# for t in range(len(NewValueList)):
#     for r in range(len(NewValueList[t])):
#         for c in range(len(NewValueList[t][r])):
#             ni=ni+1

# print(ni)

# print('ni above')






# pprint.pprint(ValueTextList)

# print(len(ValueTextList))


# print(len(ValueTextList[0]))

# print(len(ValueTextList[0][0]))


# AllValueTextListND=[]
# for t in range(len(ValueTextList)):
#     for r in range(len(ValueTextList[t])):
#         for c in range(len(ValueTextList[r])):
#             ValueTextListND = NoDuplicateList.NoDuplicateList(List=ValueTextList[r][c])
#             AllValueTextListND.append(ValueTextListND)



# pprint.pprint(AllValueTextListND)


# ValueTextListND = NoDuplicateList.NoDuplicateList(List=ValueTextList)

# print(ValueTextListND)






table=GetTableDocPositions2(filename=filename4)

pprint.pprint(table)


print("Done")