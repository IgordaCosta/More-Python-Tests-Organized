from docx import Document
import pprint




def appendTableValue(tableTextLocations,LengthOfWordToCheck,totalLengthTableWords):
    AmountOfTimesWordApeared = len(tableTextLocations)+1

    AmountOfLettersToGoBack= LengthOfWordToCheck * AmountOfTimesWordApeared

    tableTextLocations.append(totalLengthTableWords - AmountOfLettersToGoBack)





# WordTextLocations =  [55, 79, 538, 946, 974, 1927]
# tableTextLocations =  [25, 26, 27, 28, 29, 30, 31, 32, 78, 79, 80, 81, 82, 83, 84, 90, 91, 92, 99, 100, 101, 106, 146, 147, 148, 155, 156, 157, 182, 183, 184, 185, 186, 187, 188, 189, 235, 236, 237, 238, 239, 240, 241, 699, 782]
# tableTextLocationsSimplified =  [8, 7, 3, 3, 1, 3, 3, 8, 7, 1, 1]

# UserWordTextLocations = [str(tableTextLocations[i])+":_"+str(i)+"__Data_used" for i in range(len(tableTextLocations))]

# UsertableTextLocations = [str(tableTextLocationsSimplified[i])+":_"+str(i)+"__Data_used" for i in range(len(tableTextLocationsSimplified))]








# def PlaceWordInTextAndTableDataPositions2(wordDocunent,WordToCheck):
def PlaceWordInTextAndTableDataPositions2(wordDocunent,UserWordTextLocations,UsertableTextLocations,WordTextLocations,tableTextLocations, tableTextLocationsSimplified):

    # UsertableTextLocations

    # tableTextLocationsSimplified

    UsertableTextLocationsFinal = []

    for item in tableTextLocationsSimplified:
        for number in range(item):

            UsertableTextLocationsFinal.append(UsertableTextLocations[item])

    pprint.pprint(UsertableTextLocationsFinal)

    # print(len(UsertableTextLocationsFinal))

    # print(len(tableTextLocations))


    document = Document(wordDocunent)

    LengthOfWordToCheck = len(WordToCheck)

    #to get text letter locations

    WordTextLocations=[]
    totalLengthWords=0

    WordTextCount=0
    for paragraph in document.paragraphs:
        for letter in paragraph.text:
        
            totalLengthWords = totalLengthWords + 1 
            if WordTextCount == letter:

                AmountOfTimesWordApeared = len(WordTextLocations)+1

                AmountOfLettersToGoBack= LengthOfWordToCheck * AmountOfTimesWordApeared

                WordTextLocations.append(totalLengthWords + totalLengthWords - AmountOfLettersToGoBack)



    # To get table data letter locations

    lastLetter = ''
    tableTextLocations = []
    totalLengthTableWords=0
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:

                CellText=cell.text
                # print(CellText)
                
                for letter in CellText:
                
                    # print(word.text)
                    LengthOfWordToCheck = 1
                    totalLengthTableWords = totalLengthTableWords + LengthOfWordToCheck
                        
                    if WordToCheck == letter:
                        if letter == lastLetter:
                            lastLetter = ''
                            totalLengthTableWords = totalLengthTableWords - LengthOfWordToCheck
                        else:
                            print('letter: ',letter)

                            appendTableValue(tableTextLocations,LengthOfWordToCheck,totalLengthTableWords)

                    lastLetter = letter

                        
    # print(len(WordTextLocations)+len(tableTextLocations))

    return WordTextLocations, tableTextLocations







WordToCheck= 'Ÿ'

wordDocunent = r'C:\Users\IgorDC\Desktop\PydocTest\TestChanged5.docx'


# WordTextLocations =  [54, 77, 535, 942, 969, 1921]






WordTextLocations =  [55, 79, 538, 946, 974, 1927]
tableTextLocations =  [25, 26, 27, 28, 29, 30, 31, 32, 78, 79, 80, 81, 82, 83, 84, 90, 91, 92, 99, 100, 101, 106, 146, 147, 148, 155, 156, 157, 182, 183, 184, 185, 186, 187, 188, 189, 235, 236, 237, 238, 239, 240, 241, 699, 782]
tableTextLocationsSimplified =  [8, 7, 3, 3, 1, 3, 3, 8, 7, 1, 1]

UserWordTextLocations = [str(WordTextLocations[i])+":_"+str(i)+"__Data_used" for i in range(len(WordTextLocations))]

UsertableTextLocations = [str(tableTextLocationsSimplified[i])+":_"+str(i)+"__Data_used" for i in range(len(tableTextLocationsSimplified))]



print(UsertableTextLocations)


# WordTextLocations, tableTextLocations = PlaceWordInTextAndTableDataPositions2(wordDocunent,UserWordTextLocations,UsertableTextLocations,WordTextLocations,tableTextLocations, tableTextLocationsSimplified)

PlaceWordInTextAndTableDataPositions2(wordDocunent,UserWordTextLocations,UsertableTextLocations,WordTextLocations,tableTextLocations, tableTextLocationsSimplified)


# print('WordTextLocations: ', WordTextLocations )

# print('tableTextLocations: ', tableTextLocations)




# results for 'Ÿ' and 'C:\Users\IgorDC\Desktop\PydocTest\TestChanged5.docx'

# WordTextLocations:  [55, 79, 538, 946, 974, 1927]
# tableTextLocations:  [24, 69, 74, 80, 84, 123, 129, 153, 198, 655, 737]