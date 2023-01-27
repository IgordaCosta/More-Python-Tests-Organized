import docx
from docx import Document
import pprint
from io import StringIO
import pandas
import textract

import FindDocTablePositionsALL





def GetDocPositionWithVariable(filename,PositionList,LocationKey):
    # doc = docx.Document(filename)
    # f = open(filename, 'rb')
    # document = Document(f)

    # documentAll=''
    # for i in document:
    #     documentAll = documentAll.join(i)


    # print(documentAll)

    
    text = textract.process(filename)

    # print(type(text))

    # WOrd="BLLLLLLLLAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAFFFFFFFFFFFFFFFF"

    # print(text[:7]+bytes(WOrd,encoding='UTF-8'))


    docIO = StringIO(str(text))

    df = pandas.read_csv(docIO, sep="XXXXXXXXXXX", engine='python')

    # pprint.pprint(df.columns.tolist())

    print('len(df): ',len(df.columns.tolist()))

    Seprations=df.columns.tolist()

    ListOfPositions=[]
    for i in range(len(Seprations)-1):
        print('========================================================')
        print(Seprations[i])
        if i == 0:
             lengthOfSeparator=len(Seprations[i])-2
        else:
            lengthOfSeparator=len(Seprations[i])

        print('len(i):', len(Seprations[i]))
        ListOfPositions.append(lengthOfSeparator)

    print(ListOfPositions)


    
    # DefinedPositionList=[]
    # for i in range(len(PositionList)):
    #     [t, r, c ] = PositionList[i]
    #     ValueListInput = doc.tables[t].cell(r,c).text
    #     if ValueListInput==LocationKey:
    #         DefinedPositionList.append(PositionList[i])

    # return DefinedPositionList




LocationKey='XXXXXXXXXXX'

filename='C:\\Users\\IgorDC\\Desktop\\PydocTest\\TestDocNoChange.docx'

filename2='C:\\Users\\IgorDC\\Desktop\\PydocTest\\TestChanged.docx'

filename4='C:\\Users\\IgorDC\\Desktop\\PydocTest\\TestChanged4.docx'


# PositionList = FindDocTablePositionsALL.FindDocTablePositionsALL(filename)

# DefinedPositionList = GetDocPositionWithVariable(filename2,PositionList,LocationKey)

PositionList=''

GetDocPositionWithVariable(filename2,PositionList,LocationKey)

# pprint.pprint(DefinedPositionList) 

# print(len(DefinedPositionList))
