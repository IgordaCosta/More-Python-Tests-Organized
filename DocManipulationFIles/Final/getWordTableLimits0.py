######### this functions gets the length of all the columns 
######### in the word table without any empty lists and repeats



from docx import Document
import pprint




def getWordTableLimits(Location, FileName):
    
    wordDocument = Location +'\\' +FileName

    document = Document(wordDocument)

    BeforeTableLocations = []
    for t in range(document.tables.__len__()):
        docTables = document.tables[t].rows

        BeforeTableRowLocations = []
        for r in range(docTables.__len__()):
            docTableRows = docTables[r].cells

            BeforeTableCellLocations = []
            for c in range(docTableRows.__len__()):

                try:
                    BeforeAppendValue = document.tables[t].rows[r].cells[c].text
                except:
                    pass

                document.tables[t].rows[r].cells[c].text = str(c)

                BeforeTableCellLocations.append(BeforeAppendValue)

            BeforeTableRowLocations.append(BeforeTableCellLocations)

        BeforeTableLocations.append(BeforeTableRowLocations)



    TableLocations = []
    for t in range(document.tables.__len__()):
        ddocTables = document.tables[t].rows

        TableRowLocations = []
        for r in range(docTables.__len__()):
            docTableRows = docTables[r].cells

            TableCellLocations = []
            for c in range(docTableRows.__len__()):

                try:
                    AppendValue = document.tables[t].rows[r].cells[c].text
                except:
                    pass

                try:
                    AppendValue0 = document.tables[t].rows[r].cells[c - 1].text

                    if  AppendValue0 ==  AppendValue:
                        pass
                    else:
                        TableCellLocations.append(str(t) + '.'+ str(r) +'.' + AppendValue)

                except IndexError:
                    pass

            if TableCellLocations == []:
                pass
            else:
                TableRowLocations.append(TableCellLocations)
        if TableRowLocations == []:
                pass
        else:   
            TableLocations.append(TableRowLocations)

    return TableLocations   







#test bellow

Location = r'C:\Users\IgorDC\Desktop\PydocTest'

FileName = 'TestDocNoChangeAllBoldArialBlack.docx'


TableLocations = getWordTableLimits(Location, FileName)


pprint.pprint(TableLocations)


# print('=====================================================')


# pprint.pprint(BeforeTableLocations)
            

# print(len((TableLocations)))

# print(len(BeforeTableLocations))