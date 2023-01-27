from docx import Document
import pprint

import getWordTableLimitsOneList


def GetSentenceLetterLocation(sentence, y = -1, docFile=True, doc =''):
    
    for x in range(len(sentence)):
            letter = sentence[x]

            if WordToCheck == letter:
                if docFile:
                    try:
                        doc.write('letter: '+ str(letter) +"\n")
                        doc.write('Paragraphs y: '+ str(y) +"\n") 
                        doc.write('sentence x: '+ str(x) +"\n") 

                        doc.write('==============================================='+"\n") 
                    except:
                        pass
                if y == -1:
                    ParagraphSentenceLocation = str(x)
                else:
                    ParagraphSentenceLocation = str(y) + "." + str(x)


                return ParagraphSentenceLocation

            



# wordDocunent,WordToCheck, document,doc, docFile=False, ParagraphAndTableData=True
def GetTextOnlyDataPositions(wordDocunent,WordToCheck, document,doc,docFile=True, ParagraphAndTableData=True):

    if ParagraphAndTableData:
        pass
    else:   
        document = Document(wordDocunent)

    WordYXLocations = []

    for y in range(len(document.paragraphs)):
        paragraph = document.paragraphs[y]

        # print(paragraph.text)

        for x in range(len(paragraph.text)):
            letter = paragraph.text[x]

            if WordToCheck == letter:
                if docFile:
                    try:
                        doc.write('letter: '+ str(letter) +"\n")
                        doc.write('Paragraphs y: '+ str(y) +"\n") 
                        doc.write('sentence x: '+ str(x) +"\n") 

                        doc.write('==============================================='+"\n") 
                    except:
                        pass

                ParagraphSentenceLocation = str(y) + "." + str(x)

                WordYXLocations.append(ParagraphSentenceLocation)


    return WordYXLocations




# wordDocunent,WordToCheck, document,doc, docFile=False, ParagraphAndTableData=True
def GetTableDataOnlyPositions(wordDocunent,WordToCheck, document,tableLimitLocationList ,doc, docFile=True, ParagraphAndTableData=True):

    if ParagraphAndTableData:
        pass
    else:   
        document = Document(wordDocunent)

    totalLengthTableWords=0
    TableYXILocation = []

    valueToCheckInListOld = 'rr'
    # print(tableLimitLocationList)
    for y in range(len(document.tables)):
        table = document.tables[y]
        for x in range(len(table.rows)):
            row= table.rows[x]
            for i in range(len(row.cells)):

               
                valueToCheckInList = str(y)+"."+str(x)+"."+str(i)
                # print(valueToCheckInList)

                # print(valueToCheckInList in tableLimitLocationList)
                if valueToCheckInList in tableLimitLocationList:
                   
                    RowTextLength = len(row.cells[i].text)

                    NotAddedYet = True
                    for letter in range(RowTextLength):
                    
                        LengthOfWordToCheck = 1
                        
                        totalLengthTableWords = totalLengthTableWords + LengthOfWordToCheck

                        LengthOfCellText = str(len(row.cells[i].text))

                        

                        if LengthOfCellText == '1':
                            ToAddWordYXLocations=''

                        else:

                            DocumentLocationNotText = row.cells[i]

                            WordYXLocations = GetTextOnlyDataPositions(wordDocunent,WordToCheck, document = DocumentLocationNotText, ParagraphAndTableData=True, docFile=True, doc = doc)
                            
                            WordYXLocations0 = str(WordYXLocations)

                            WordYXLocations00=WordYXLocations0.replace(" ","")

                            WordYXLocations1 = WordYXLocations00.replace(",",".")
                            WordYXLocations2 = WordYXLocations1.replace("'","")
                            WordYXLocations3 = WordYXLocations2.replace('"','')
                            WordYXLocations4 =  WordYXLocations3.replace('[','')
                            WordYXLocations5 =  WordYXLocations4.replace(']','')

                            if WordYXLocations5 == "":
                                ToAddWordYXLocations = WordYXLocations5
                            else:
                                ToAddWordYXLocations = "." + WordYXLocations5

                            # print(ToAddWordYXLocations)

                        if docFile:
                            try:
                                doc.write('Table y: '+ str(y) +"\n") 
                                doc.write('Row x: '+ str(x) +"\n") 
                                doc.write('Row Cell [i]: '+ str(i) +"\n") 
                                doc.write('totalLengthTableWords: '+str(totalLengthTableWords) +"\n") 
                                doc.write('row.cells[i].text: '+ row.cells[i].text +"\n") 
                                doc.write('LengthOfCellText: '+ LengthOfCellText +"\n") 
                                doc.write('letter in number: '+ str(letter)+"/"+str(RowTextLength-1) +"\n")
                                doc.write('letter: '+ str(row.cells[i].text[letter]) +"\n")
                                doc.write('up till letter: '+ str(row.cells[i].text[:letter +1]) +"\n")
                                
                                doc.write('============================================='+"\n")
                            except:
                                pass

                        # print('WordToCheck == str(row.cells[i].text[letter]): ' + str(WordToCheck == str(row.cells[i].text[letter])))
                        # print('str(row.cells[i].text[letter] ' + str(row.cells[i].text[letter]))
                        if WordToCheck == str(row.cells[i].text[letter]):
                            
                            if NotAddedYet:
                                valueTOAppend = str(y)+"."+str(x)+"."+str(i)+"."+str(letter)+"."+LengthOfCellText + ToAddWordYXLocations
                                TableYXILocation.append(valueTOAppend)

                                NotAddedYet = False #just added

    

    return TableYXILocation




def GetTextAndTableDataPosition(wordDocunent,WordToCheck, docFile=True, docSaveFileName=''):

    document = Document(wordDocunent)
    
    if docFile:
        doc = open(r"C:\Users\IgorDC\Desktop\PydocTest\testFile.txt","w")#write mode 
    else:
        doc = ''


    tableLimitLocationList = getWordTableLimitsOneList.getWordTableLimitsOneList(FileName=wordDocunent)


    # To get table data locations
    TableYXILocation = GetTableDataOnlyPositions(wordDocunent,WordToCheck, document, tableLimitLocationList, doc, docFile=False, ParagraphAndTableData=True)


    # To get text letter locations
    WordYXLocations = GetTextOnlyDataPositions(wordDocunent,WordToCheck, document,doc, docFile=False , ParagraphAndTableData=True)

    
    
    if docFile:
        doc.close()

    return WordYXLocations , TableYXILocation

    








WordToCheck= 'Ÿ'

# WordToCheck= ' Ÿ'

# wordDocunent = r'C:\Users\IgorDC\Desktop\PydocTest\TestChanged5.docx'

# wordDocunent = r'C:\Users\IgorDC\Desktop\PydocTest\TestChanged12.docx'

wordDocunent = r'C:\Users\IgorDC\Desktop\PydocTest\TestChanged13.docx'

# wordDocunent = r'C:\Users\IgorDC\Desktop\PydocTest\DifferentStylesTestChangedTNR16_3.docx'


# wordDocunent = r'C:\Users\IgorDC\Desktop\PydocTest\DifferentStylesTestChangedTNR16_3_WithCode.docx'

# wordDocunent = r'C:\Users\IgorDC\Desktop\PydocTest\DifferentStylesTestChangedTNR16_3_WithCode2.docx'

docSaveFileName = r"C:\Users\IgorDC\Desktop\PydocTest\testFile.txt"

docFile = False






# WordTextLocations, tableTextLocations, tableTextLocationsSimplified = GetTextAndTableDataPosition(wordDocunent,WordToCheck)


WordYXLocations , TableYXILocation = GetTextAndTableDataPosition(wordDocunent = wordDocunent, WordToCheck = WordToCheck, docFile = docFile, docSaveFileName = docSaveFileName)


# GetTextAndTableDataPosition(wordDocunent = wordDocunent, WordToCheck = WordToCheck, docFile = docFile, docSaveFileName = docSaveFileName)


pprint.pprint('WordYXLocations: ') 

pprint.pprint(WordYXLocations)


pprint.pprint('TableYXILocation: ')

# pprint.pprint(TableYXILocation)

print(TableYXILocation)




# results for 'Ÿ' and 'C:\Users\IgorDC\Desktop\PydocTest\TestChanged5.docx'

# WordTextLocations:  [55, 79, 538, 946, 974, 1927]

# tableTextLocations:  [25, 26, 27, 28, 29, 30, 31, 32, 78, 79, 80, 81, 82, 83, 84, 90, 91, 92, 99, 100, 101, 106, 146, 147, 148, 155, 156, 157, 182, 183, 184, 185, 186, 187, 188, 189, 235, 236, 237, 238, 239, 240, 241, 699, 782]
# tableTextLocationsSimplified:  [8, 7, 3, 3, 1, 3, 3, 8, 7, 1, 1]