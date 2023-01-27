from docx import Document

import UserImputToTableImput
import closeWord







def appendTableValue(tableTextLocations,LengthOfWordToCheck,totalLengthTableWords):
    AmountOfTimesWordApeared = len(tableTextLocations)+1

    AmountOfLettersToGoBack= LengthOfWordToCheck * AmountOfTimesWordApeared

    tableTextLocations.append(totalLengthTableWords - AmountOfLettersToGoBack)



# WordTextLocations:  [55, 79, 538, 946, 974, 1927]
# tableTextLocations:  [24, 69, 74, 80, 84, 123, 129, 153, 198, 655, 737]
# AddedAt_ 

#  _XXXXXXXXXXXXXXXXXXXXXX
# WordTextLocations:  [54, 77, 535, 942, 969, 1921]
# WordTextLocationsToAddList = ['AddedAt_54_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_77_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_535_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_942_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_969_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_1921_XXXXXXXXXXXXXXXXXXXXX']

# def GetTextAndTableDataPosition(wordDocunent,WordToCheck):
# def GetTextAndTableDataPosition(wordDocunent,WordTextLocations,tableTextLocations,AddSideSpaces=True):
# def GetTextAndTableDataPosition(wordDocument,WordTextLocations,WordTextLocationsToAddList, TableTextLocations, TableTextLocationsToAddList,AddSideSpaces=True):
def PlaceWordInTextAndTableDataPositions(wordDocument,WordTextLocations,WordTextLocationsToAddList, TableTextLocations, tableTextLocationsSimplified, UserTableImput, AddSideSpaces=True):


    doc = open(r"C:\Users\IgorDC\Desktop\PydocTest\testFile.txt","w")#write mode 


    TableTextLocationsToAddList = UserImputToTableImput.UserImputToTableImput(tableTextLocationsSimplified, UserImput=UserTableImput)

    document = Document(wordDocument)

    AtWordStep = -1

    DocumentLegth=0
    # for i in range(len(document.paragraphs)):
       
    #     paragraphLength=len(document.paragraphs[i].text)

    #     DocumentLegthOld= DocumentLegth

    #     DocumentLegth = DocumentLegth + paragraphLength

    #     AllLenWord = 0
    #     insideParagraph = True
    #     numberOfRunsInside = 0
    #     while insideParagraph:
    #         paragraphLength=len(document.paragraphs[i].text)

    #         try:
    #             AtWordLocation = WordTextLocations[AtWordStep +1]


    #             if DocumentLegth >= AtWordLocation:
    #                 if  AtWordLocation >= DocumentLegthOld:

    #                     AtWordStep = AtWordStep + 1

    #                     LocationToAddData=AtWordLocation-DocumentLegthOld

    #                     TotalSpaceToForward = LocationToAddData 

    #                     WordToAdd=WordTextLocationsToAddList[AtWordStep]

    #                     if AddSideSpaces:
    #                         WordToAdd= " "+ WordToAdd + " "
                        
    #                     document.paragraphs[i].text = document.paragraphs[i].text[:TotalSpaceToForward + AllLenWord -1 ] +WordToAdd + document.paragraphs[i].text[TotalSpaceToForward + AllLenWord: ]

    #                     LenWord= len(WordToAdd)

    #                     AllLenWord= AllLenWord + LenWord -1

    #                     numberOfRunsInside= numberOfRunsInside + 1
                    
    #                 else:
    #                     insideParagraph = False
                
    #             else:
    #                 insideParagraph = False
                

    #         except IndexError:
    #             insideParagraph = False




    # print(document.tables[4].rows[2].cells[9].text)


    # document.tables[4].rows[2].cells[9].text = "PPPPPPPPPPPPPPPPPPPPP"

    # lastLetter = ''
    # tableTextLocations = []
    # totalLengthTableWords=0
    # TableTextLocationsIndex = 0
    # MaxTableTextLocationsIndex = len(TableTextLocations)
    # LengthOfWordToCheck = 1

    # itenToAdd = 'Ð'

    # SubLength = 0
    # for y in range(len(document.tables)):
    #     table = document.tables[y]
    #     for x in range(len(table.rows)):
    #         row= table.rows[x]
    #         for i in range(len(row.cells)):
                
    #             letterRealIndex = 0

    #             RowTextLength = len(row.cells[i].text)
                
    #             for letter in range(RowTextLength):
                
    #                 LengthOfWordToCheck = 1
                    
    #                 totalLengthTableWords = totalLengthTableWords + LengthOfWordToCheck


    #                 doc.write('Table y: '+ str(y) +"\n") 
    #                 doc.write('Row x: '+ str(x) +"\n") 
    #                 doc.write('Row Cell [i]: '+ str(i) +"\n") 
    #                 doc.write('totalLengthTableWords: '+str(totalLengthTableWords) +"\n") 
    #                 doc.write('row.cells[i].text: '+ row.cells[i].text +"\n") 
    #                 doc.write('letter in number: '+ str(letter)+"/"+str(RowTextLength-1) +"\n")
    #                 doc.write('letter: '+ str(row.cells[i].text[letter]) +"\n")
    #                 doc.write('up till letter: '+ str(row.cells[i].text[:letter +1]) +"\n")
                    
    #                 doc.write('============================================='+"\n")

    #                 if TableTextLocationsIndex < MaxTableTextLocationsIndex:

    #                     TableValue = TableTextLocations[TableTextLocationsIndex]

    #                     TableValue = TableValue + SubLength
    #                     # doc.write('y: '+ str(y) +"\n") 
    #                     # doc.write('x: '+ str(x) +"\n") 
    #                     # doc.write('[i]: '+ str(i) +"\n") 
    #                     # doc.write('totalLengthTableWords: '+str(totalLengthTableWords) +"\n") 
    #                     # doc.write('row.cells[i].text: '+ row.cells[i].text +"\n") 
    #                     # doc.write('============================================='+"\n")

    #                     # print('totalLengthTableWords: ',totalLengthTableWords)

    #                     # print('row.cells[i].text: ', row.cells[i].text)

    #                     if TableValue == totalLengthTableWords:

    #                         # print('row.cells[i].text: ', row.cells[i].text)

    #                         # OrigianlLength = len(row.cells[i].text)


    #                         # # row.cells[i].text = row.cells[i].text[:letterRealIndex] + itenToAdd + row.cells[i].text[letterRealIndex+1:]
    #                         # row.cells[i].text = row.cells[i].text[:letterRealIndex] + itenToAdd + row.cells[i].text[letterRealIndex:]
              

    #                         # AfterLength = len(row.cells[i].text)

    #                         # SubLength = AfterLength - OrigianlLength

    #                         TableTextLocationsIndex = TableTextLocationsIndex +1

                        
    #                     letterRealIndex = letterRealIndex + 1

                            
    # for table in document.tables:
    #     for row in table.rows:

    #         for i in range(len(row.cells)):
    
    #             for letter in range(len(row.cells[i].text)):

    #                 if row.cells[i].text[letter] == itenToAdd:
 
    #                     row.cells[i].text = row.cells[i].text[letter+1:] + "QQQQQ" + row.cells[i].text[:letter]

                
                
                # row.cells[i].text = FinalCellLook




                        
                    # if WordToCheck == letter:
                    #     if letter == lastLetter:
                    #         lastLetter = ''
                    #         totalLengthTableWords = totalLengthTableWords - LengthOfWordToCheck
                    #     else:
                    #         print('letter: ',letter)

                    #         appendTableValue(tableTextLocations,LengthOfWordToCheck,totalLengthTableWords)

                    # lastLetter = letter
        


    # AtWordStep = -1
    # DocumentLegth=0
    # for table in document.tables:
    #     for row in table.rows:

    #         for i in range(len(row.cells)):

    #             for letter in range(len(row.cells[i].text)):

    #                 if row.cells[i].text[letter] == 'Ÿ':
 
    #                     row.cells[i].text = row.cells[i].text[letter+1:] + "QQQQQ" + row.cells[i].text[:letter]



    fileName = r'C:\Users\IgorDC\Desktop\PydocTest\ChangedWordDocument20210130.docx'                 
    try:
        document.save(fileName)
    except PermissionError:
        closeWord.closeWord(fileName=fileName)
        document.save(fileName)

    doc.close()
    print("Done")


    # document.save(r'C:\Users\IgorDC\Desktop\PydocTest\ChangedWordDocument20210130.docx')





                # for letter in CellText:
                #     print(letter)
                
                #     # print(word.text)
                #     LengthOfWordToCheck = 1
                #     totalLengthTableWords = totalLengthTableWords + LengthOfWordToCheck
                        
                #     if WordToCheck == letter:
                #         if letter == lastLetter:
                #             lastLetter = ''
                #             totalLengthTableWords = totalLengthTableWords - LengthOfWordToCheck
                #         else:
                #             # print('letter: ',letter)
                #             appendTableValue(tableTextLocations,LengthOfWordToCheck,totalLengthTableWords)

                #     lastLetter = letter

                        
    # print(len(WordTextLocations)+len(tableTextLocations))

    # return WordTextLocations, tableTextLocations




WordTextLocations =  [55, 79, 538, 946, 974, 1927]
tableTextLocations =  [25, 26, 27, 28, 29, 30, 31, 32, 78, 79, 80, 81, 82, 83, 84, 90, 91, 92, 99, 100, 101, 106, 146, 147, 148, 155, 156, 157, 182, 183, 184, 185, 186, 187, 188, 189, 235, 236, 237, 238, 239, 240, 241, 699, 782]
tableTextLocationsSimplified =  [8, 7, 3, 3, 1, 3, 3, 8, 7, 1, 1]


WordTextLocationsToAddList = ['AddedAt_54_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_77_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_535_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_942_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_969_XXXXXXXXXXXXXXXXXXXXX', 'AddedAt_1921_XXXXXXXXXXXXXXXXXXXXX']


TableTextLocations =  [25, 26, 27, 28, 29, 30, 31, 32, 78, 79, 80, 81, 82, 83, 84, 90, 91, 92, 99, 100, 101, 106, 146, 147, 148, 155, 156, 157, 182, 183, 184, 185, 186, 187, 188, 189, 235, 236, 237, 238, 239, 240, 241, 699, 782]


UserTableImput = ['AddedAt_:'+str(tableTextLocationsSimplified[i])+"__:"+str(i)+"_XXXXXXXXXXXXXXXXXXXXX" for i in range(len(tableTextLocationsSimplified))]


WordToCheck= 'Ÿ'

# wordDocument = r'C:\Users\IgorDC\Desktop\PydocTest\TestDocNoChange.docx'

wordDocument = r'C:\Users\IgorDC\Desktop\PydocTest\TestChanged11.docx'


PlaceWordInTextAndTableDataPositions(wordDocument,WordTextLocations,WordTextLocationsToAddList, TableTextLocations, tableTextLocationsSimplified, UserTableImput, AddSideSpaces=True)



# WordTextLocations, tableTextLocations = GetTextAndTableDataPosition(wordDocunent,WordToCheck)


# print('WordTextLocations: ', WordTextLocations )

# print('tableTextLocations: ', tableTextLocations)




# results for 'Ÿ' and 'C:\Users\IgorDC\Desktop\PydocTest\TestChanged5.docx'

# WordTextLocations:  [54, 77, 535, 942, 969, 1921]
# tableTextLocations:  [24, 69, 74, 80, 84, 123, 129, 153, 198, 655, 737]