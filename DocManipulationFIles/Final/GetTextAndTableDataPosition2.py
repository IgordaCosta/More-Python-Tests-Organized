from docx import Document



def GetTextAndTableDataPosition(wordDocunent,WordToCheck):

    document = Document(wordDocunent)

    # doc = open(r"C:\Users\IgorDC\Desktop\PydocTest\testFile.txt","w")#write mode 

    #to get text letter locations
    WordYXLocations=[]
    # totalLengthWords=0
    for y in range(len(document.paragraphs)):
        paragraph = document.paragraphs[y]
        for x in range(len(paragraph.text)):
            letter = paragraph.text[x]
            # totalLengthWords = totalLengthWords + 1 

            if WordToCheck == letter:
                # doc.write('letter: '+ str(letter) +"\n")
                # doc.write('Paragraphs y: '+ str(y) +"\n") 
                # doc.write('sentence x: '+ str(x) +"\n") 
                # doc.write('totalLengthWords: '+ str(totalLengthWords) +"\n")
                # doc.write('==============================================='+"\n") 

                ParagraphSentenceLocation = str(y) + "." + str(x)

                WordYXLocations.append(ParagraphSentenceLocation)


    # print(WordYXLocations)


    # To get table data letter locations
    # totalLengthTableWords=0
    TableYXILocation = []
    for y in range(len(document.tables)):
        table = document.tables[y]
        for x in range(len(table.rows)):
            row= table.rows[x]
            for i in range(len(row.cells)):

                RowTextLength = len(row.cells[i].text)
                
                for letter in range(RowTextLength):
                
                    # LengthOfWordToCheck = 1
                    
                    # totalLengthTableWords = totalLengthTableWords + LengthOfWordToCheck


                    # doc.write('Table y: '+ str(y) +"\n") 
                    # doc.write('Row x: '+ str(x) +"\n") 
                    # doc.write('Row Cell [i]: '+ str(i) +"\n") 
                    # doc.write('totalLengthTableWords: '+str(totalLengthTableWords) +"\n") 
                    # doc.write('row.cells[i].text: '+ row.cells[i].text +"\n") 
                    # doc.write('letter in number: '+ str(letter)+"/"+str(RowTextLength-1) +"\n")
                    # doc.write('letter: '+ str(row.cells[i].text[letter]) +"\n")
                    # doc.write('up till letter: '+ str(row.cells[i].text[:letter +1]) +"\n")
                    
                    # doc.write('============================================='+"\n")

                    
                    if WordToCheck == str(row.cells[i].text[letter]):
                        # print(" check if "+WordToCheck+" = "+str(row.cells[i].text[letter]))
                        valueTOAppend = str(y)+"."+str(x)+"."+str(i)
                        TableYXILocation.append(valueTOAppend)


    # print(TableYXILocation)

    # doc.close()

    return WordYXLocations , TableYXILocation

    








WordToCheck= 'Ÿ'

# wordDocunent = r'C:\Users\IgorDC\Desktop\PydocTest\TestChanged5.docx'

wordDocunent = r'C:\Users\IgorDC\Desktop\PydocTest\TestChanged12.docx'




# WordTextLocations, tableTextLocations, tableTextLocationsSimplified = GetTextAndTableDataPosition(wordDocunent,WordToCheck)


WordYXLocations , TableYXILocation = GetTextAndTableDataPosition(wordDocunent,WordToCheck)


print('WordYXLocations: ', WordYXLocations)


print('TableYXILocation: ', TableYXILocation)




# results for 'Ÿ' and 'C:\Users\IgorDC\Desktop\PydocTest\TestChanged5.docx'

# WordTextLocations:  [55, 79, 538, 946, 974, 1927]

# tableTextLocations:  [25, 26, 27, 28, 29, 30, 31, 32, 78, 79, 80, 81, 82, 83, 84, 90, 91, 92, 99, 100, 101, 106, 146, 147, 148, 155, 156, 157, 182, 183, 184, 185, 186, 187, 188, 189, 235, 236, 237, 238, 239, 240, 241, 699, 782]
# tableTextLocationsSimplified:  [8, 7, 3, 3, 1, 3, 3, 8, 7, 1, 1]