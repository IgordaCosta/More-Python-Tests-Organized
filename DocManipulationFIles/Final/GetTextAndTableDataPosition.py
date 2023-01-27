from docx import Document



def GetTextAndTableDataPosition(wordDocunent,WordToCheck):

    document = Document(wordDocunent)

    LengthOfWordToCheck = len(WordToCheck)

    #to get text letter locations
    WordTextLocations=[]
    totalLengthWords=0
    for paragraph in document.paragraphs:
        for letter in paragraph.text:

            totalLengthWords = totalLengthWords + 1 

            if WordToCheck == letter:

                WordTextLocations.append(totalLengthWords)


    # To get table data letter locations
    tableTextLocations = []
    tableTextLocationsSimplified = []
    totalLengthTableWords=0

    Cellsize = 1
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:

                CellText=cell.text
                # print(CellText)
                
                for letter in CellText:
                
                    # print(word.text)
                    LengthOfWordToCheck = 1
                    totalLengthTableWords = totalLengthTableWords + LengthOfWordToCheck
                    
                    if WordToCheck == letter:

                        sizeOftableTextLocations = len(tableTextLocations)

                        if sizeOftableTextLocations == 0:

                            tableTextLocations.append(totalLengthTableWords)

                        else:

                            oldtableTextLocationsValue = tableTextLocations[-1]
                            
                            tableTextLocations.append(totalLengthTableWords)

                            if oldtableTextLocationsValue + 1 == totalLengthTableWords:
                                Cellsize = Cellsize + 1

                            else:
                                tableTextLocationsSimplified.append(Cellsize)
                                # print('Cellsize: ',Cellsize)
                                Cellsize = 1

                            # print('Cellsize: ',Cellsize)
                            # print('tableTextLocations: ',totalLengthTableWords)


            

    #THIS IS THE LAST VALUE TO APPEND
    tableTextLocationsSimplified.append(Cellsize)



    return WordTextLocations, tableTextLocations, tableTextLocationsSimplified







WordToCheck= 'Ÿ'

wordDocunent = r'C:\Users\IgorDC\Desktop\PydocTest\TestChanged5.docx'




WordTextLocations, tableTextLocations, tableTextLocationsSimplified = GetTextAndTableDataPosition(wordDocunent,WordToCheck)


print('WordTextLocations: ', WordTextLocations )

print('tableTextLocations: ', tableTextLocations)

print("tableTextLocationsSimplified: ", tableTextLocationsSimplified)




# results for 'Ÿ' and 'C:\Users\IgorDC\Desktop\PydocTest\TestChanged5.docx'

# WordTextLocations:  [55, 79, 538, 946, 974, 1927]

# tableTextLocations:  [25, 26, 27, 28, 29, 30, 31, 32, 78, 79, 80, 81, 82, 83, 84, 90, 91, 92, 99, 100, 101, 106, 146, 147, 148, 155, 156, 157, 182, 183, 184, 185, 186, 187, 188, 189, 235, 236, 237, 238, 239, 240, 241, 699, 782]
# tableTextLocationsSimplified:  [8, 7, 3, 3, 1, 3, 3, 8, 7, 1, 1]