######### this functions gets the length of all the columns 
######### in the word table without any empty lists and repeats



from docx import Document
import pprint




def getWordTableLimitsOneList(Location = '', FileName = '',document=''):
    
    goOn = False

    if document == '':
        wordDocument = Location +'\\' +FileName
       
        if FileName == '':
            print('add data')
            pass
        elif Location =='':
            if FileName == '':
                print('addData')
            else: 
                wordDocument = FileName
                document = Document(wordDocument)
                goOn=True
        else:
            document = Document(wordDocument)
            goOn=True
    
    else:
        goOn=True
    
    if goOn:

        t = -1 
        BeforeTableLocations = []
        for docTables in document.tables:
            
            t = t +1
            
            r = -1
            BeforeTableRowLocations = []
            for docTableRows in docTables.rows:
                
                r = r +1
            
                c = -1
                BeforeTableCellLocations = []
                for Item in docTableRows.cells:
                    
                    c = c +1

                    Item.text = str(c)



        
        t = -1
        TableLocations = []
        for docTables in document.tables:
            
            t = t + 1
            
            r = -1
            TableRowLocations = []
            for docTableRows in docTables.rows:

                TableCellLocations = []
                
                AppendValue0 ='rr'
                okToGO= False
                r = r + 1
                for Item in docTableRows.cells:

                    try:
                        AppendValue = Item.text
                        okToGO= True
                    except:
                        pass

                    if okToGO:

                        if  AppendValue0 ==  AppendValue:
                            pass
                        else:
                            TableCellLocations.append(str(t) + '.'+ str(r) +'.' + AppendValue)
                            AppendValue0 = AppendValue

                    

                if TableCellLocations == []:
                    pass
                else:
                    TableRowLocations.append(TableCellLocations)
            if TableRowLocations == []:
                    pass
            else:   
                TableLocations.append(TableRowLocations)

        
        SingleListTableLocations = []
        for itemX in TableLocations:
            for itemY in itemX:
                for item in itemY:
                    SingleListTableLocations.append(item)
                    
        
        return SingleListTableLocations
        
        
        # return TableLocations   







#test bellow

# Location = r'C:\Users\IgorDC\Desktop\PydocTest'

# FileName = 'TestDocNoChangeAllBoldArialBlack.docx'


# TableLocations = getWordTableLimitsOneList(Location=Location, FileName=FileName)


# pprint.pprint(TableLocations)


