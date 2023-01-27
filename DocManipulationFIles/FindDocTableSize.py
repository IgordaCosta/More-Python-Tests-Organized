import docx
from docx import Document


filename='C:\\Users\\IgorDC\\Desktop\\PydocTest\\TestDocNoChange.docx'


splitter='XXXXXXXXXXX'

filename2='C:\\Users\\IgorDC\\Desktop\\PydocTest\\TestChanged.docx'


filename4='C:\\Users\\IgorDC\\Desktop\\PydocTest\\TestChanged4.docx'





wordDoc = Document(filename2)


indexNumber=0
TableRowCell=[]
indexList=[]


# print(len(wordDoc.tables))

# print('len(wordDoc.tables)')



# add digit values from 1 to amount to end into all table variables
# then get those values back, the result will be be variables with empty spots in the middle
#  ex. 1,4,8,12,20,25 ...

tableNumber=0
for table in wordDoc.tables:

    # print('len(wordDoc.tables): '+ str(len(wordDoc.tables)))
    

    # print('tableNumber: ' + str(tableNumber))

    
    rowNumber=0
    for row in table.rows:

        # print("len(table.rows): " + str(len(table.rows)))
        
        # print('rowNumber: ' + str(rowNumber))

        
        

        cellNumber=0
        id=0
        for cell in row.cells:

            # print('len(row.cells) :' + str(len(row.cells)))
            
            # print('cellNumber: ' + str(cellNumber))


            wordDoc.tables[tableNumber].cell(rowNumber,cellNumber).text = str(id)

            id = id+1

            
            
            cellNumber=cellNumber+1

        rowNumber=rowNumber+1

    tableNumber=tableNumber+1




wordDoc.save(filename4)


print("Done")